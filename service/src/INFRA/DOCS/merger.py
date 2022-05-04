import json
import os
from docx import Document
from docxcompose.composer import Composer
from docx.shared import Inches
from docx.shared import Pt
from datetime import datetime
from docx2pdf import convert
from datetime import timedelta

FILE_DIR = os.path.dirname(__file__)
ROOT_DIR = os.path.dirname(FILE_DIR)
DOC_DIR = os.path.join(ROOT_DIR, ("web\\static\\docs"))
TMP_DIR = os.path.join(FILE_DIR, (f"tmp\\"))
DOCS = []
ARRIVAL = datetime.now()
DEPARTURE = ARRIVAL + timedelta(days=1)
WEEK = ["Monday", "Tuesday",
        "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"]


def gen_tour_doc(data):
    try:
        header = gen_header(data)
        travel_code = header.get("travel_code")
        gen_cover_doc(header, travel_code)
        gen_dest(data, travel_code)
        doc_compose(DOCS, travel_code)
    except Exception as e:
        print(e)
        raise e


def gen_header(data):
    try:
        customer_data = data.get("customer")
        tour_data = data.get("tour")
        logistic_data = data.get("logistic")
        header = {**customer_data, **tour_data, **logistic_data}
        return header
    except Exception as e:
        print(e)
        raise e


def replace_word(document, key, word):
    try:
        for paragraph in document.paragraphs:
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, word)
    except Exception as e:
        print(e)
        raise e


def check_date(date, format):
    try:
        return datetime.strptime(date, format)
    except:
        return date


def gen_cover_doc(header, name):
    try:
        passengers = header.get("passengers")
        arrival_date = header.get("arrival_date")
        departure_date = header.get("departure_date")
        until_date = header.get("until_date")
        names = header.get("names").upper()
        last_names = header.get("last_names").upper()
        ARRIVAL = check_date(arrival_date, "%Y-%m-%d %H:%M:%S")
        DEPARTURE = check_date(departure_date, "%Y-%m-%d %H:%M:%S")
        until = check_date(until_date, "%Y-%m-%d %H:%M:%S.%f")
        days = (DEPARTURE - ARRIVAL).days
        nights = days - 1
        customer = f"{names} {last_names}"
        print(header)
        composed = os.path.join(DOC_DIR, (f"{name}.docx"))
        validate_docx(composed)
        doc_template = os.path.join(FILE_DIR, (f"cover.docx"))
        doc_output = os.path.join(TMP_DIR, (f"cover-{name}.docx"))
        document = Document(doc_template)
        replace_word(document, "NIGHTS", f"{nights}")
        replace_word(document, "DAYS", f"{days}")
        replace_word(document, "CUSTOMER", f"{customer.upper()}")
        replace_word(document, "LEGAL_NAME", f"{customer.upper()}")
        replace_word(document, "PAX", f"{passengers}")
        replace_word(document, "VALID", f'{until.strftime("%Y-%m-%d")}')
        document.save(doc_output)
        DOCS.append(doc_output)
    except Exception as e:
        print(e)
        raise e


def gen_dest(data, name):
    try:
        destinations = data.get("destinations")
        for dest_id, dest_name in enumerate(destinations):
            destination = destinations[dest_name]
            dest_name = destination["destination"]
            doc = Document()
            run = doc.add_heading(f"{dest_name.upper()}", 0).add_run()
            font = run.font
            font.name = "Calibri"
            font.size = Pt(12)
            days = destination.get("daysData")
            for day_id, day_name in enumerate(days):
                day = days[day_name]
                experiences = day.get("experiences")
                meals = day.get("meals")
                type = "Tour"
                date = (ARRIVAL + timedelta(days=int(day_name))
                        )
                date_string = date.strftime("%Y-%m-%d")
                weekday = WEEK[date.weekday()]
                if (day_id) == 0:
                    type = f"Arrival"
                    doc.add_heading(f" {date_string} {weekday}", 1)
                    meals = "D/O"
                elif day_id == len(days) - 1 and dest_id == len(destinations) - 1:
                    type = "Departure"
                    doc.add_heading(f" {date_string} {weekday}", 1)
                    meals = "B/L"
                else:
                    type = "Tour"
                    doc.add_heading(
                        f"Day {day_name} {date_string} {weekday}", 1)
                doc.add_heading(f"Experiences", 2)
                for i, exp_name in enumerate(experiences):
                    experience = experiences[exp_name]
                    run = doc.add_heading(exp_name, 3).add_run()
                    font = run.font
                    font.name = "Calibri"
                    font.size = Pt(12)
                    description = experience.get("description").split(".")[1:]
                    run = doc.add_paragraph(
                        description
                    ).add_run()
                    font = run.font
                    font.name = "Calibri"
                    font.size = Pt(12)
                    image = os.path.join(
                        FILE_DIR, (f"images\\{dest_name}.png"))
                    doc.add_picture(image, width=Inches(4), height=Inches(4))
                    if i < len(experiences) - 1:
                        next = list(experiences)[i + 1]
                        next_description = experiences[next]['description'].split('.')[0]
                        run = doc.add_paragraph(
                            f"Next we're going to get, {next_description}"
                        ).add_run()
                        font = run.font
                        font.name = "Calibri"
                        font.size = Pt(12)
                    else:
                        if type == "Departure":
                            run = doc.add_paragraph(
                                f"Take any photos to go back to Home"
                            ).add_run()
                        else:
                            run = doc.add_paragraph(
                                f"Next we're going back to hotel to rest and take a meal\n"
                            ).add_run()
                        font = run.font
                        font.name = "Calibri"
                        font.size = Pt(12)
                        run = doc.add_paragraph(f"{meals}").add_run()
                        font = run.font
                        font.name = "Calibri"
                        font.size = Pt(12)
                    # else:
                    #     doc.add_paragraph(f"Next we going to departure to Home", style='Intense Quote')
                # doc.add_page_break()
            output = os.path.join(
                TMP_DIR, (f"{dest_name}-{name}-{dest_id}.docx"))
            doc.save(output)
            DOCS.append(output)
    except Exception as e:
        print(e)
        raise e

def validate_docx(file):
    try:
        if os.path.exists(file):
            os.remove(file)
    except Exception as e:
        return False

def doc_compose(files, name):
    try:
        composed = os.path.join(DOC_DIR, (f"{name}.docx"))
        validate_docx(composed)
        pdf = os.path.join(DOC_DIR, (f"{name}.pdf"))
        new_document = Document()
        composer = Composer(new_document)
        for i in range(0, len(files)):
            doc = Document(files[i])
            if i != len(files) - 1:
                doc.add_page_break()
            composer.append(doc)
        composer.save(composed)
        # convert(composed, pdf)
    except Exception as e:
        print(e)
        raise e


def empty_tmp():
    try:
        from pathlib import Path

        for f in Path(TMP_DIR).glob('*.*'):
            try:
                f.unlink()
            except OSError as e:
                print("Error: %s : %s" % (f, e.strerror))
    except Exception as e:
        print(e)
        raise e


# if __name__ == "__main__":
#     empty_tmp()
#     DATA_PATH = os.path.join(FILE_DIR, ("data.json"))
#     test_data = open(DATA_PATH, "r", encoding="utf8")
#     data = json.load(test_data)
#     gen_tour_doc(data)
#     empty_tmp()
